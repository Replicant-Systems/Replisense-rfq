# file_parser.py
from pathlib import Path
import pandas as pd
import docx
import asyncio
import aiofiles
from typing import Dict, Any, Optional
import logging
from concurrent.futures import ThreadPoolExecutor
import json
from functools import lru_cache
import hashlib

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    import PyPDF2
    HAS_PDFPLUMBER = False
    logging.warning("pdfplumber not available, falling back to PyPDF2. Install pdfplumber for better PDF parsing.")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileParsingError(Exception):
    """Custom exception for file parsing errors"""
    pass

class FileParser:
    def __init__(self, max_file_size_mb: int = 10, max_excel_rows: int = 1000):
        self.max_file_size = max_file_size_mb * 1024 * 1024  # Convert to bytes
        self.max_excel_rows = max_excel_rows
        self.thread_pool = ThreadPoolExecutor(max_workers=4)
        
        # Supported file types
        self.supported_extensions = {'.txt', '.pdf', '.xlsx', '.xls', '.docx', '.csv', '.json'}
        
        logger.info(f"FileParser initialized with max_file_size={max_file_size_mb}MB, max_excel_rows={max_excel_rows}")

    async def parse_file_async(self, file_path: str) -> Dict[str, Any]:
        """Async version of parse_file"""
        file_path = Path(file_path)
        
        # Validate file
        await self._validate_file_async(file_path)
        
        # Get file hash for caching
        file_hash = await self._get_file_hash_async(file_path)
        
        # Check cache first
        cached_result = self._get_cached_result(file_hash)
        if cached_result:
            logger.info(f"Using cached result for {file_path.name}")
            return cached_result
        
        # Parse file based on extension
        try:
            if file_path.suffix.lower() == '.txt':
                text = await self._parse_txt_async(file_path)
            elif file_path.suffix.lower() == '.pdf':
                text = await self._parse_pdf_async(file_path)
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                text = await self._parse_excel_async(file_path)
            elif file_path.suffix.lower() == '.docx':
                text = await self._parse_docx_async(file_path)
            elif file_path.suffix.lower() == '.csv':
                text = await self._parse_csv_async(file_path)
            elif file_path.suffix.lower() == '.json':
                text = await self._parse_json_async(file_path)
            else:
                raise FileParsingError(f"Unsupported file format: {file_path.suffix}")
            
            result = {
                "raw_text": text,
                "source_file": str(file_path.name),
                "file_size": file_path.stat().st_size,
                "file_hash": file_hash,
                "parsing_method": f"async_{file_path.suffix.lower()[1:]}"
            }
            
            # Cache result
            self._cache_result(file_hash, result)
            
            logger.info(f"Successfully parsed {file_path.name} ({len(text)} characters)")
            return result
            
        except Exception as e:
            logger.error(f"Error parsing {file_path.name}: {str(e)}")
            raise FileParsingError(f"Failed to parse {file_path.name}: {str(e)}")

    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """Synchronous version for backward compatibility"""
        return asyncio.run(self.parse_file_async(file_path))

    async def _validate_file_async(self, file_path: Path) -> None:
        """Validate file exists, size, and type"""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise FileParsingError(f"File too large: {file_size / 1024 / 1024:.1f}MB (max: {self.max_file_size / 1024 / 1024}MB)")
        
        # Check file extension
        if file_path.suffix.lower() not in self.supported_extensions:
            raise FileParsingError(f"Unsupported file type: {file_path.suffix}")
        
        # Check if file is readable
        try:
            async with aiofiles.open(file_path, 'rb') as f:
                await f.read(1)  # Try to read first byte
        except PermissionError:
            raise FileParsingError(f"Permission denied reading file: {file_path}")

    async def _get_file_hash_async(self, file_path: Path) -> str:
        """Generate hash for file caching"""
        hash_md5 = hashlib.md5()
        async with aiofiles.open(file_path, 'rb') as f:
            chunk = await f.read(8192)
            while chunk:
                hash_md5.update(chunk)
                chunk = await f.read(8192)
        return hash_md5.hexdigest()

    @lru_cache(maxsize=50)
    def _get_cached_result(self, file_hash: str) -> Optional[Dict[str, Any]]:
        """Get cached parsing result"""
        # In a real implementation, you might use Redis or database
        return None

    def _cache_result(self, file_hash: str, result: Dict[str, Any]) -> None:
        """Cache parsing result"""
        # In a real implementation, you might use Redis or database
        pass

    async def _parse_txt_async(self, file_path: Path) -> str:
        """Parse text file asynchronously"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = await f.read()
                return content.strip()
        except UnicodeDecodeError:
            # Fallback to different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        content = await f.read()
                        logger.info(f"Successfully read {file_path.name} using {encoding} encoding")
                        return content.strip()
                except:
                    continue
            raise FileParsingError("Unable to decode text file with any supported encoding")

    async def _parse_pdf_async(self, file_path: Path) -> str:
        """Parse PDF file asynchronously with improved extraction"""
        if HAS_PDFPLUMBER:
            return await asyncio.get_event_loop().run_in_executor(
                self.thread_pool, self._parse_pdf_with_pdfplumber, file_path
            )
        else:
            return await asyncio.get_event_loop().run_in_executor(
                self.thread_pool, self._parse_pdf_with_pypdf2, file_path
            )

    def _parse_pdf_with_pdfplumber(self, file_path: Path) -> str:
        """Parse PDF using pdfplumber (better for tables and complex layouts)"""
        import pdfplumber
        
        text_parts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"=== Page {page_num} ===\n{page_text}")
                    
                    # Extract tables if any
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table:
                            table_text = f"\n=== Table {table_num} on Page {page_num} ===\n"
                            for row in table:
                                if row:
                                    table_text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                            text_parts.append(table_text)
                        
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path.name}: {str(e)}")
            # Fallback to PyPDF2
            return self._parse_pdf_with_pypdf2(file_path)
        
        result = "\n\n".join(text_parts)
        return result if result.strip() else "Unable to extract text from PDF"

    def _parse_pdf_with_pypdf2(self, file_path: Path) -> str:
        """Fallback PDF parsing with PyPDF2"""
        import PyPDF2
        
        text_parts = []
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(f"=== Page {page_num} ===\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num} from {file_path.name}: {str(e)}")
                        continue
        except Exception as e:
            raise FileParsingError(f"PDF parsing failed: {str(e)}")
        
        result = "\n\n".join(text_parts)
        return result if result.strip() else "Unable to extract text from PDF"

    async def _parse_excel_async(self, file_path: Path) -> str:
        """Parse Excel file asynchronously with row limits"""
        return await asyncio.get_event_loop().run_in_executor(
            self.thread_pool, self._parse_excel_sync, file_path
        )

    def _parse_excel_sync(self, file_path: Path) -> str:
        """Parse Excel file with improved handling"""
        text_parts = []
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None, nrows=self.max_excel_rows)
            
            for sheet_name, df in excel_data.items():
                if df.empty:
                    continue
                
                sheet_text = f"=== Sheet: {sheet_name} ===\n"
                
                # Handle large sheets
                if len(df) > self.max_excel_rows:
                    sheet_text += f"Note: Showing first {self.max_excel_rows} rows of {len(df)} total rows\n"
                    df = df.head(self.max_excel_rows)
                
                # Clean and format data
                df = df.dropna(how='all')  # Remove completely empty rows
                df = df.fillna('')  # Replace NaN with empty string
                
                # Convert to readable format
                sheet_text += df.to_string(index=False, max_rows=None)
                text_parts.append(sheet_text)
                
        except Exception as e:
            raise FileParsingError(f"Excel parsing failed: {str(e)}")
        
        if not text_parts:
            return "Excel file appears to be empty or unreadable"
        
        return "\n\n".join(text_parts)

    async def _parse_docx_async(self, file_path: Path) -> str:
        """Parse DOCX file asynchronously with table support"""
        return await asyncio.get_event_loop().run_in_executor(
            self.thread_pool, self._parse_docx_sync, file_path
        )

    def _parse_docx_sync(self, file_path: Path) -> str:
        """Parse DOCX file with improved content extraction"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            # Extract paragraphs
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if paragraphs:
                content_parts.append("=== Document Text ===")
                content_parts.extend(paragraphs)
            
            # Extract tables
            if doc.tables:
                content_parts.append("\n=== Tables ===")
                for table_num, table in enumerate(doc.tables, 1):
                    content_parts.append(f"\n--- Table {table_num} ---")
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            content_parts.append(row_text)
            
            result = "\n".join(content_parts)
            return result if result.strip() else "DOCX file appears to be empty"
            
        except Exception as e:
            raise FileParsingError(f"DOCX parsing failed: {str(e)}")

    async def _parse_csv_async(self, file_path: Path) -> str:
        """Parse CSV file asynchronously"""
        return await asyncio.get_event_loop().run_in_executor(
            self.thread_pool, self._parse_csv_sync, file_path
        )

    def _parse_csv_sync(self, file_path: Path) -> str:
        """Parse CSV file with improved handling"""
        try:
            # Try to read CSV with different encodings and separators
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                for sep in [',', ';', '\t']:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, sep=sep, nrows=self.max_excel_rows)
                        if len(df.columns) > 1:  # Valid CSV with multiple columns
                            df = df.dropna(how='all').fillna('')
                            result = f"=== CSV Data (using {encoding}, separator '{sep}') ===\n"
                            result += df.to_string(index=False, max_rows=None)
                            return result
                    except Exception:
                        continue
            
            raise FileParsingError("Unable to parse CSV with any supported encoding/separator combination")
            
        except Exception as e:
            raise FileParsingError(f"CSV parsing failed: {str(e)}")

    async def _parse_json_async(self, file_path: Path) -> str:
        """Parse JSON file asynchronously"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
                
            # Validate JSON
            data = json.loads(content)
            
            # Format for better readability
            formatted_json = json.dumps(data, indent=2, ensure_ascii=False)
            return f"=== JSON Data ===\n{formatted_json}"
            
        except json.JSONDecodeError as e:
            raise FileParsingError(f"Invalid JSON format: {str(e)}")
        except Exception as e:
            raise FileParsingError(f"JSON parsing failed: {str(e)}")

    def __del__(self):
        """Cleanup thread pool"""
        if hasattr(self, 'thread_pool'):
            self.thread_pool.shutdown(wait=False)


# Convenience functions
async def parse_file_async(file_path: str, max_file_size_mb: int = 10) -> Dict[str, Any]:
    """Convenience function for async file parsing"""
    parser = FileParser(max_file_size_mb=max_file_size_mb)
    return await parser.parse_file_async(file_path)

def get_supported_extensions() -> set:
    """Get list of supported file extensions"""
    return {'.txt', '.pdf', '.xlsx', '.xls', '.docx', '.csv', '.json'}